from flask import Flask, render_template, request, redirect, jsonify
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from datetime import datetime
from docx.shared import Pt
from docx import Document
import urllib.parse
import pandas as pd
import threading
import shutil
import json
import os
import re

load_dotenv()
app = Flask(__name__)

excel_folder = os.getenv('EXCEL_FOLDER')
doc_template = os.getenv('DOC_TEMPLATE')
doc_result = os.getenv('DOC_RESULT')

ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}

processing_status = {
    'total_files': 0,
    'processed_files': 0,
    'status': 'Waiting for file upload'
}
folder_to_delete = None


def logs_to_json(error_from_function, folder_name, error):
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    data_to_store = {
        "from": error_from_function,
        "folder_name": folder_name,
        "time": current_time,
        "error": error
    }

    try:
        with open(folders('storage', 'logs.json'), 'r+') as file:
            stored_data = json.load(file)
            stored_data.append(data_to_store)
            file.seek(0)
            json.dump(stored_data, file, indent=4)
    except FileNotFoundError:
        with open(folders('storage', 'logs.json'), 'w') as file:
            json.dump([data_to_store], file, indent=4)


def folders(path, file_name, project_name=None, user_name=None):
    try:
        if project_name is None:
            folder_path = path
        else:
            if user_name is None:
                folder_path = os.path.join(path, project_name)
            else:
                folder_path = os.path.join(path, project_name, user_name)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        return os.path.join(folder_path, file_name)
    except Exception as e:
        print(e)


def delete_excel_folder(folder_to_delete):
    try:
        excel_path = os.path.join(os.getcwd(), excel_folder, folder_to_delete)
        if os.path.exists(excel_path):
            shutil.rmtree(excel_path)
            return "Excel folder deleted successfully"
        else:
            return "Excel folder does not exists"
    except Exception as e:
        logs_to_json('delete_excel_folder_function', folder_to_delete, str(e))
        return "Error while deleting excel folder"


def allowed_file(filename):
    try:
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    except Exception as e:
        logs_to_json('allowed_file', 'allowed_file', str(e))


def replace_text(doc, placeholder, replacement):
    try:
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(replacement))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text(cell, placeholder, replacement)
    except Exception as e:
        logs_to_json('replace_text', 'replace_text', str(e))


def set_formatting(doc):
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.highlight_color = None

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.font.highlight_color = None
    except Exception as e:
        logs_to_json('set_formatting', 'set_formatting', str(e))


def process_excel(file_path, folder_name):
    try:
        global processing_status
        file_extension = file_path.rsplit('.', 1)[1].lower()
        df_zalogoderzhatel = None
        df_zalogodatel = None

        if file_extension == 'csv':
            df_main = pd.read_csv(file_path)
        else:
            df_main = pd.read_excel(file_path, sheet_name='текст ДЗ заполнен')
            df_main.columns = df_main.columns.str.strip()
            df_zalogodatel = pd.read_excel(file_path, sheet_name='Залогодатель', header=None)
            df_zalogoderzhatel = pd.read_excel(file_path, sheet_name='Залогодержатель', header=None)

            df_zalogoderzhatel = df_zalogoderzhatel.T
            df_zalogoderzhatel.columns = df_zalogoderzhatel.iloc[0]
            df_zalogoderzhatel = df_zalogoderzhatel.drop(df_zalogoderzhatel.index[0]).reset_index(drop=True)
            df_zalogoderzhatel.columns = df_zalogoderzhatel.columns.str.strip()

            df_zalogodatel = df_zalogodatel.T
            df_zalogodatel.columns = df_zalogodatel.iloc[0]
            df_zalogodatel = df_zalogodatel.drop(df_zalogodatel.index[0]).reset_index(drop=True)
            df_zalogodatel.columns = df_zalogodatel.columns.str.strip()

        output_files = []
        processing_status['total_files'] = len(df_main)

        for index, row in df_main.iterrows():
            doc = Document('doc_template/Проект ДКП.docx')
            contract_string = df_main.iloc[index]['Договор купли продажи']
            match = re.search(r'№\s*([^от]+?)\s+от', contract_string)
            contract_number = match.group(1).strip() if match else None
            date_match = re.search(r'от\s*(\d{2})\.(\d{2})\.(\d{4})', contract_string)

            if contract_number is not None:
                replace_text(doc, 'P_section', contract_number)
            if date_match is not None:
                replace_text(doc, 'dd_p', date_match.group(1))
                replace_text(doc, 'mm_p', date_match.group(2))
                replace_text(doc, 'yy_p', date_match.group(3))

            replace_text(doc, 'ТОО_продавец', df_zalogoderzhatel['Наименование компании'].iloc[0])
            replace_text(doc, 'должность_продавца', df_zalogoderzhatel['в лице (Подписанта) — должности'].iloc[0])
            replace_text(doc, 'ФИО_продавца', df_zalogoderzhatel['в лице (Подписанта) - Фамилии И.О.'].iloc[0])
            replace_text(doc, 'ИИН_продавца', df_zalogoderzhatel['БИН'].iloc[0])
            replace_text(doc, 'должность_покупателя', df_zalogodatel['в лице (Подписанта) — должности'].iloc[0])
            replace_text(doc, 'ФИО_покупателя', df_zalogodatel['в лице (Подписанта) - Фамилии И.О.'].iloc[0])
            replace_text(doc, 'ИИН_покупателя', df_zalogodatel['БИН'].iloc[0])
            replace_text(doc, 'общая_площадь', df_main.iloc[index]['Общая площадь'])
            replace_text(doc, 'дом_номер, кв. квартира_номер,', df_main.iloc[index]['Адрес квартиры'])
            replace_text(doc, 'кадастровый_номер', df_main.iloc[index]['РКА'])
            replace_text(doc, 'на_основании_чего от', df_main.iloc[index]['№ договора ПДКП/дата'])
            replace_text(doc, 'сумма_тенге', df_main.iloc[index]['Сумма уступленного долга по ПДКП (тенге)'])
            date_value = df_main.iloc[index]['Срок полного выкупа (ПДКП)']
            formatted_date = pd.to_datetime(date_value).strftime('%Y-%m-%d')
            replace_text(doc, 'год_расчета', formatted_date)
            replace_text(doc, 'БИН_продавца', df_zalogoderzhatel['БИН'].iloc[0])
            replace_text(doc, 'ИИК_продавца', df_zalogoderzhatel['IBAN'].iloc[0])
            replace_text(doc, 'БИК_продавца', df_zalogoderzhatel['БИК'].iloc[0])
            replace_text(doc, 'Банк_продаца', df_zalogoderzhatel['Банк'].iloc[0])
            replace_text(doc, 'телефон_продавца', df_zalogoderzhatel['Контактные телефоны:'].iloc[0])
            replace_text(doc, 'подпись_должность_пр', df_zalogoderzhatel['Подписант, должность'].iloc[0])
            replace_text(doc, 'подпись_должность_пок', df_zalogodatel['Подписант, должность'].iloc[0])

            set_formatting(doc)

            output_filename = str(folders(doc_result, f'Проект ДКП_{index + 1}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx', folder_name))
            doc.save(output_filename)
            output_files.append(output_filename)

            # Update progress
            processing_status['processed_files'] = index + 1

        processing_status['status'] = 'Договоры созданы успешно!' if len(output_files) > 0 else 'Error'
        return output_files
    except Exception as e:
        logs_to_json('process_excel', 'process_excel', str(e))


@app.route('/', methods=['GET', 'POST'])
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        global processing_status
        global folder_to_delete
        if 'file' not in request.files:
            return redirect(request.url)

        file = request.files['file']

        if file.filename == '':
            return redirect(request.url)

        if file and allowed_file(file.filename):
            name, _ = os.path.splitext(file.filename)
            folder_to_delete = name
            encoded_filename = urllib.parse.quote(file.filename)
            filename = secure_filename(encoded_filename)
            file_path = str(folders(excel_folder, filename, name))
            file.save(file_path)

            processing_status['status'] = 'Processing started'
            processing_status['total_files'] = 0
            processing_status['processed_files'] = 0
            threading.Thread(target=process_excel, args=(file_path, name)).start()

            return redirect('/status')

        return redirect(request.url)
    except Exception as e:
        logs_to_json('upload_file', 'upload_file', str(e))


@app.route('/status')
def status():
    try:
        if processing_status['status'] == 'Договоры созданы успешно!':
            delete_excel_folder(folder_to_delete)
        return jsonify(processing_status)
    except Exception as e:
        logs_to_json('status', 'status', str(e))


if __name__ == '__main__':
    app.run()
