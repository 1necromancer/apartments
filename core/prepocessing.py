from layout.baseUtils import logs_to_json, folders
from dotenv import load_dotenv
from datetime import datetime
from docx.shared import Pt
from docx import Document
from pathlib import Path
import pandas as pd
import os

ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}
processing_status = {
    'total_files': 0,
    'processed_files': 0,
    'status': 'Waiting for file upload'
}
env_path = Path(__file__).resolve().parent.parent / '.env'
load_dotenv(dotenv_path=env_path)
doc_result = os.getenv('DOC_RESULT')
doc_template = os.getenv('DOC_TEMPLATE')


def allowed_file(filename):
    try:
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    except Exception as e:
        logs_to_json('allowed_file', 'allowed_file', str(e))


def replace_text(doc, placeholder, replacement):
    try:
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(replacement))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text(cell, placeholder, replacement)
    except Exception as e:
        logs_to_json('replace_text', 'replace_text', str(e))


def set_formatting(doc):
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.highlight_color = None

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.font.highlight_color = None
    except Exception as e:
        logs_to_json('set_formatting', 'set_formatting', str(e))


def process_excel(file_path, folder_name, template_name):
    try:
        global processing_status
        file_extension = file_path.rsplit('.', 1)[1].lower()
        dict_mappings = {}
        sheet_data = {}

        if file_extension == 'csv':
            df_main = pd.read_csv(file_path)
        else:
            sheet_names = ['главная', 'покупатель', 'продавец']

            for sheet in sheet_names:
                if sheet in ['покупатель', 'продавец']:
                    df_sheet = pd.read_excel(file_path, sheet_name=sheet, header=None)
                    df_sheet = df_sheet.T
                    df_sheet.columns = df_sheet.iloc[0]
                    df_sheet = df_sheet.drop(df_sheet.index[0])
                    df_sheet.reset_index(drop=True, inplace=True)
                else:
                    df_sheet = pd.read_excel(file_path, sheet_name=sheet)

                df_sheet.columns = df_sheet.columns.str.strip()
                sheet_data[sheet] = df_sheet

            df_dictionary = pd.read_excel(file_path, sheet_name='справочник')
            df_dictionary.columns = df_dictionary.columns.str.strip()

            df_main = sheet_data['главная']

            for _, row in df_dictionary.iterrows():
                change_from = row['поменять с'].strip()
                change_to = row['поменять на'].strip()
                sheet_name = row['название листа'].strip()

                if sheet_name not in dict_mappings:
                    dict_mappings[sheet_name] = {}
                dict_mappings[sheet_name][change_from] = change_to

        output_files = []
        processing_status['total_files'] = len(df_main)
        change_to_value = None

        for index, row in df_main.iterrows():
            doc = Document(f'{doc_template}/{template_name}.docx')

            for sheet, replacements in dict_mappings.items():
                if sheet in sheet_data:
                    df_sheet = sheet_data[sheet]

                    for change_from, change_to_column in replacements.items():
                        try:
                            if sheet == 'главная':
                                if index < len(df_sheet):
                                    change_to_value = df_sheet[change_to_column].iloc[index]
                            else:
                                change_to_value = df_sheet[change_to_column].iloc[0]

                            replace_text(doc, change_from, str(change_to_value))
                        except Exception as e:
                            logs_to_json('process_excel', 'perhaps column does not exist', str(e))

            set_formatting(doc)
            output_filename = folders(doc_result,
                                      f'{template_name}_{index + 1}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx',
                                      folder_name)
            doc.save(output_filename)
            output_files.append(output_filename)

            processing_status['processed_files'] = index + 1

        processing_status['status'] = 'Договоры созданы успешно!' if len(output_files) > 0 else 'Error'
        return output_files

    except Exception as e:
        logs_to_json('process_excel', 'process_excel', str(e))
